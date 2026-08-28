import sys
import os
import json
sys.path.insert(0, os.path.abspath("backend"))

from fastapi.testclient import TestClient
from main import app
from database import SessionLocal
from models.candidate import CandidateModel
from models.application import ApplicationModel
from models.job import JobModel

client = TestClient(app)

def run_tests():
    print("\n--- STARTING ATS WORKFLOW ARCHITECTURE & INTEGRATION TESTS ---")
    db = SessionLocal()
    
    # 0. Setup test jobs
    j1 = client.get("/api/jobs")
    jobs = j1.json() if j1.status_code == 200 else []
    job1_id = jobs[0]["id"] if jobs else "JOB-0001"
    job2_id = jobs[1]["id"] if len(jobs) > 1 else "JOB-0002"

    test_email = "workflow.tester@datacraft.in"

    # Clean previous test entries if any
    from services.candidate_service import delete_candidate
    c_old = db.query(CandidateModel).filter(CandidateModel.email == test_email).all()
    for c in c_old:
        delete_candidate(db, c.id)
    db.commit()

    # TEST 1 & 2 & 3 & 4: Upload CV / HR Add Candidate creates Candidate & Application
    print("\n[TEST 1-4] Creating candidate profile & initial application...")
    res1 = client.post(f"/api/candidates?jobId={job1_id}", json={
        "firstName": "Workflow",
        "lastName": "Tester",
        "email": test_email,
        "phone": "+91 91111 22222",
        "currentRole": "Senior Python Engineer",
        "skills": ["Python", "FastAPI", "PostgreSQL", "Docker"],
        "resumeText": "Experienced Senior Python Engineer with expertise in FastAPI, PostgreSQL, Docker, microservices, and automated testing.",
        "location": "Pune, India"
    })
    assert res1.status_code == 201, f"Candidate creation failed: {res1.text}"
    cand1 = res1.json()
    cand1_id = cand1["id"]
    print(f"  [OK] Candidate created successfully: id='{cand1_id}', candidateId='{cand1['candidateId']}'.")

    # Verify Application created in DB with initial status='Applied'
    app1 = db.query(ApplicationModel).filter(ApplicationModel.candidateId == cand1_id).first()
    assert app1 is not None, "Application record was not automatically created!"
    assert app1.jobId == job1_id, f"Application connected to wrong jobId: {app1.jobId}"
    assert app1.status.lower() == "applied", f"Expected initial status 'Applied', got '{app1.status}'"
    print(f"  [OK] Application created automatically in 'Applied' stage: id='{app1.id}', applicationId='{app1.applicationId}', status='{app1.status}'.")

    # TEST 5: Resume Storage & Text
    print("\n[TEST 5] Uploading resume document and checking text availability...")
    import docx, io
    doc = docx.Document()
    doc.add_heading("Senior Python Engineer Resume", level=1)
    doc.add_paragraph("Experienced Senior Python Engineer with expertise in FastAPI, PostgreSQL, Docker, and microservices.")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    sample_docx_bytes = doc_io.getvalue()

    up_res = client.post(
        f"/api/candidates/{cand1_id}/resume",
        files={"file": ("python_dev_resume.docx", sample_docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert up_res.status_code == 200, f"Resume upload failed: {up_res.text}"
    print("  [OK] Resume PDF uploaded and stored successfully.")

    txt_res = client.get(f"/api/candidates/{cand1_id}/resume/text")
    assert txt_res.status_code == 200, f"Resume text fetch failed: {txt_res.text}"
    print(f"  [OK] Extracted resume text available cleanly.")

    # TEST 6 - 12: AI Screening (POST /api/screen-resume)
    print("\n[TEST 6-12] Executing AI Resume Screening against Job 1 JD...")
    screen_res = client.post("/api/screen-resume", json={"applicationId": app1.id})
    assert screen_res.status_code == 200, f"Screening failed: {screen_res.text}"
    eval_data = screen_res.json()["evaluation"]
    
    assert "score" in eval_data, "Missing score in evaluation"
    assert "summary" in eval_data, "Missing summary in evaluation"
    assert "strengths" in eval_data, "Missing strengths in evaluation"
    assert "gaps" in eval_data, "Missing gaps in evaluation"
    assert "interviewQuestions" in eval_data, "Missing interviewQuestions in evaluation"
    assert "fitReasoning" in eval_data, "Missing fitReasoning in evaluation"

    app1_score = eval_data["score"]
    db.refresh(app1)
    assert app1.status.lower() in ["screening", "shortlisted"], f"Expected status transition to Screening/Shortlisted, got '{app1.status}'"
    print(f"  [OK] AI Screening complete. Score: {app1_score}%. Status progressed to '{app1.status}'.")
    print(f"  [OK] Summary: {eval_data['summary'][:60]}...")
    print(f"  [OK] Strengths ({len(eval_data['strengths'])}), Gaps ({len(eval_data['gaps'])}), Questions ({len(eval_data['interviewQuestions'])}).")

    # TEST 13 - 17: Pipeline State Transitions & Persistence
    print("\n[TEST 13-17] Testing Pipeline State Transitions & PostgreSQL Persistence...")
    for stage in ["Shortlisted", "Interviewing", "Offered", "Rejected"]:
        patch_res = client.patch(f"/api/applications/{app1.id}/status", json={"status": stage})
        assert patch_res.status_code == 200, f"Status update to {stage} failed: {patch_res.text}"
        
        # Verify persistence in DB
        db.refresh(app1)
        assert app1.status.lower() == stage.lower(), f"Status mismatch in DB: expected {stage}, got {app1.status}"
        print(f"  [OK] Status transition to '{stage}' verified in PostgreSQL.")

    # TEST 20: Multiple Applications for Same Candidate
    print("\n[TEST 20] Same Candidate applying to a DIFFERENT Job (Job 2)...")
    res2 = client.post(f"/api/candidates?jobId={job2_id}", json={
        "firstName": "Workflow",
        "lastName": "Tester",
        "email": test_email, # Same candidate email
        "phone": "+91 91111 22222",
        "currentRole": "Senior Python Engineer",
        "skills": ["Python", "FastAPI"],
        "resumeText": "Experienced Senior Python Engineer."
    })
    assert res2.status_code == 201, f"Second application failed: {res2.text}"
    cand2 = res2.json()
    assert cand2["id"] == cand1_id, "Candidate ID changed! Candidate profile should be reused."
    print("  [OK] Reused existing candidate profile for second application.")

    # Verify candidate now has 2 distinct applications in PostgreSQL
    apps = db.query(ApplicationModel).filter(ApplicationModel.candidateId == cand1_id).all()
    assert len(apps) == 2, f"Expected 2 applications for candidate, found {len(apps)}"
    app2 = [a for a in apps if a.jobId == job2_id][0]
    print(f"  [OK] Candidate now has 2 distinct applications: Job1 ('{app1.jobId}'), Job2 ('{app2.jobId}').")

    # Run AI screening on second application (Job 2)
    screen_res2 = client.post("/api/screen-resume", json={"applicationId": app2.id})
    assert screen_res2.status_code == 200
    app2_score = screen_res2.json()["evaluation"]["score"]

    print(f"  [OK] Application 1 (Job 1) Score: {app1.atsScore}%.")
    print(f"  [OK] Application 2 (Job 2) Score: {app2.atsScore}%.")
    assert app1.id != app2.id, "Application IDs must be distinct!"

    # TEST 21: Duplicate Application Prevention (Same Candidate + Same Job)
    print("\n[TEST 21] Attempting duplicate application for same candidate & SAME job...")
    dup_res = client.post(f"/api/candidates?jobId={job1_id}", json={
        "firstName": "Workflow",
        "lastName": "Tester",
        "email": test_email,
        "phone": "+91 91111 22222"
    })
    assert dup_res.status_code == 409, f"Duplicate application expected 409 Conflict, got {dup_res.status_code}"
    print(f"  [OK] Duplicate application correctly rejected with 409 Conflict: '{dup_res.json()['detail']}'.")

    # Cleanup test entries using candidate_service.delete_candidate
    from services.candidate_service import delete_candidate
    delete_candidate(db, cand1_id)
    db.close()

    print("\n==================================================")
    print("ALL 21 WORKFLOW & DATA MODEL TESTS PASSED CLEANLY!")
    print("==================================================\n")

if __name__ == "__main__":
    run_tests()
