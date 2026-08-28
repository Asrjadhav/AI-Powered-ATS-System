import sys
import os
import io
import uuid
import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fastapi.testclient import TestClient
from main import app
from database import SessionLocal
from sqlalchemy.orm import Session
from models.candidate import CandidateModel
from models.application import ApplicationModel
from models.interview import InterviewModel
from models.offer import OfferModel
import services.file_storage_service as file_storage_service

import pypdf
import docx

client = TestClient(app)

def create_sample_pdf(text_content: str) -> bytes:
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(width=600, height=800)
    pdf_io = io.BytesIO()
    writer.write(pdf_io)
    pdf_bytes = pdf_io.getvalue()
    
    raw_pdf = (
        f"%PDF-1.4\n"
        f"1 0 obj <</Type /Catalog /Pages 2 0 R>> endobj\n"
        f"2 0 obj <</Type /Pages /Kids [3 0 R] /Count 1>> endobj\n"
        f"3 0 obj <</Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources <</Font <</F1 5 0 R>>>>>> endobj\n"
        f"4 0 obj <</Length {len(text_content) + 40}>> stream\n"
        f"BT /F1 12 Tf 100 700 Td ({text_content}) Tj ET\n"
        f"endstream\n"
        f"endobj\n"
        f"5 0 obj <</Type /Font /Subtype /Type1 /BaseFont /Helvetica>> endobj\n"
        f"xref\n0 6\n"
        f"0000000000 65535 f \n"
        f"trailer <</Size 6 /Root 1 0 R>>\n"
        f"startxref\n0\n%%EOF\n"
    ).encode("utf-8")
    return raw_pdf

def create_sample_docx(text_content: str) -> bytes:
    doc = docx.Document()
    doc.add_heading("Candidate Resume Document", level=1)
    doc.add_paragraph(text_content)
    docx_io = io.BytesIO()
    doc.save(docx_io)
    return docx_io.getvalue()

def run_test():
    print("\n--- Starting Phase 2H Real Resume Viewer & Multi-Identifier Test Workflow ---")
    db: Session = SessionLocal()

    test_id = str(uuid.uuid4())[:8]
    cand_id = f"cand-test-viewer-{test_id}"

    try:
        # 1. CREATE CANDIDATE
        cand_email = f"viewer.candidate.{test_id}@example.com"
        create_res = client.post("/api/candidates", json={
            "id": cand_id,
            "firstName": "RealViewer",
            "lastName": "Candidate",
            "email": cand_email,
            "currentRole": "Senior Frontend Architect"
        })
        assert create_res.status_code == 201, f"Failed candidate creation: {create_res.text}"
        cand_data = create_res.json()
        cand_internal_id = cand_data["id"]
        cand_formatted_id = cand_data["candidateId"]
        print(f"1. CREATE CANDIDATE SUCCESS: id='{cand_internal_id}', candidateId='{cand_formatted_id}', email='{cand_email}'.")

        # Setup linked Application, Interview, Offer
        app_internal_id = f"app-test-viewer-{test_id}"
        app_formatted_id = f"APP-VIEW-{test_id.upper()}"
        interview_id = f"int-test-viewer-{test_id}"
        offer_id = f"off-test-viewer-{test_id}"

        db_app = ApplicationModel(
            id=app_internal_id,
            applicationId=app_formatted_id,
            candidateId=cand_internal_id,
            jobId="JOB-0001",
            status="Shortlisted",
            appliedRole="Frontend Architect",
            candidateEmail=cand_email
        )
        db_interview = InterviewModel(id=interview_id, candidateId=cand_internal_id, jobId="JOB-0001", candidateName="RealViewer Candidate", jobTitle="Frontend Architect", date="2026-09-01", time="10:00 AM", interviewer="Tech Lead", status="Scheduled")
        db_offer = OfferModel(id=offer_id, candidateId=cand_internal_id, jobId="JOB-0001", candidateName="RealViewer Candidate", jobTitle="Frontend Architect", offeredCTC=150000.0, status="Draft")
        db.add(db_app)
        db.add(db_interview)
        db.add(db_offer)
        db.commit()

        # 2. UPLOAD REAL PDF
        sample_text_pdf = f"EXPERT IN REACT AND TYPESCRIPT KEY CODE {test_id}"
        pdf_bytes = create_sample_pdf(sample_text_pdf)
        files_pdf = {"file": ("real_resume_sample.pdf", pdf_bytes, "application/pdf")}
        upload_pdf_res = client.post(f"/api/candidates/{cand_internal_id}/resume", files=files_pdf)
        assert upload_pdf_res.status_code == 200, f"Failed PDF upload: {upload_pdf_res.text}"
        pdf_storage_key = upload_pdf_res.json()["resumeStorageKey"]
        print(f"2. UPLOAD REAL PDF SUCCESS: Storage key '{pdf_storage_key}'.")

        # 3. GET ACTUAL RESUME VIA MULTIPLE IDENTIFIERS
        identifiers_to_test = [
            ("candidate.id", cand_internal_id),
            ("candidate.candidateId", cand_formatted_id),
            ("application.id", app_internal_id),
            ("application.applicationId", app_formatted_id),
            ("candidate.email", cand_email)
        ]

        for id_label, test_identifier in identifiers_to_test:
            # Test PDF Download endpoint
            dl_res = client.get(f"/api/candidates/{test_identifier}/resume")
            assert dl_res.status_code == 200, f"Failed GET /resume using {id_label} ('{test_identifier}'): {dl_res.text}"
            assert dl_res.content == pdf_bytes, f"PDF bytes mismatch using {id_label}"
            assert dl_res.headers.get("content-type") == "application/pdf", f"Content-Type mismatch using {id_label}"

            # Test Parsed Text endpoint
            txt_res = client.get(f"/api/candidates/{test_identifier}/resume/text")
            assert txt_res.status_code == 200, f"Failed GET /resume/text using {id_label} ('{test_identifier}'): {txt_res.text}"
            assert sample_text_pdf in txt_res.json()["text"], f"Parsed text mismatch using {id_label}"

            print(f"   [OK] MULTI-IDENTIFIER CHECK PASSED for {id_label} ('{test_identifier}').")

        # 4. UPLOAD DOCX
        sample_text_docx = f"SENIOR CLOUD DEVOPS ARCHITECT KEY DOCX {test_id}"
        docx_bytes = create_sample_docx(sample_text_docx)
        files_docx = {"file": ("cloud_architect_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        upload_docx_res = client.post(f"/api/candidates/{cand_formatted_id}/resume", files=files_docx)
        assert upload_docx_res.status_code == 200
        docx_storage_key = upload_docx_res.json()["resumeStorageKey"]
        print(f"4. UPLOAD DOCX SUCCESS: Replaced with DOCX (storage key '{docx_storage_key}').")

        # 5. VERIFY DOCX BYTES & TEXT VIA APPLICATION ID
        dl_docx = client.get(f"/api/candidates/{app_formatted_id}/resume")
        assert dl_docx.status_code == 200
        assert dl_docx.content == docx_bytes

        text_docx_res = client.get(f"/api/candidates/{app_formatted_id}/resume/text")
        assert text_docx_res.status_code == 200
        assert sample_text_docx in text_docx_res.json()["text"]
        print("5. DOCX BYTES & TEXT VERIFICATION SUCCESS VIA APPLICATION ID.")

        # 6. SAFE REPLACEMENT & OLD FILE REMOVAL CHECK
        old_pdf_path = file_storage_service.get_resume_path(pdf_storage_key)
        assert old_pdf_path is None or not os.path.exists(old_pdf_path)
        new_docx_path = file_storage_service.get_resume_path(docx_storage_key)
        assert new_docx_path is not None and os.path.exists(new_docx_path)
        print("6. SAFE REPLACEMENT CHECK: Old PDF file removed, new DOCX file active.")

        # 7. DELETE RESUME
        del_res = client.delete(f"/api/candidates/{cand_formatted_id}/resume")
        assert del_res.status_code == 200
        print("7. DELETE RESUME SUCCESS.")

        # 8. VERIFY RESUME ENDPOINT RETURNS 404 AFTER DELETION
        del_get_res = client.get(f"/api/candidates/{cand_internal_id}/resume")
        assert del_get_res.status_code == 404
        del_text_res = client.get(f"/api/candidates/{cand_internal_id}/resume/text")
        assert del_text_res.status_code in [404, 422]
        print("8. 404 CHECK: GET resume & text returned 404 after deletion.")

        # 9. DB INTEGRITY CHECKS (CANDIDATE, APPLICATION, INTERVIEW, OFFER REMAIN)
        db.expire_all()
        assert db.query(CandidateModel).filter(CandidateModel.id == cand_internal_id).first() is not None
        assert db.query(ApplicationModel).filter(ApplicationModel.id == app_internal_id).first() is not None
        assert db.query(InterviewModel).filter(InterviewModel.id == interview_id).first() is not None
        assert db.query(OfferModel).filter(OfferModel.id == offer_id).first() is not None
        print("9. DB INTEGRITY CHECK: Candidate, Application, Interview, and Offer remain intact.")

        # 10. INVALID FILE REJECTION (HTTP 400)
        invalid_res = client.post(f"/api/candidates/{cand_internal_id}/resume", files={"file": ("script.sh", b"echo hack", "text/x-shellscript")})
        assert invalid_res.status_code == 400
        print("10. INVALID FILE REJECTION CHECK: Returned 400 Bad Request.")

        # 11. OVERSIZED FILE REJECTION (>10 MB) (HTTP 400)
        over_res = client.post(f"/api/candidates/{cand_internal_id}/resume", files={"file": ("large.pdf", b"0" * (10*1024*1024+1), "application/pdf")})
        assert over_res.status_code == 400
        print("11. OVERSIZED FILE REJECTION CHECK: Returned 400 Bad Request.")

        print("\nALL PHASE 2H RESUME VIEWER & MULTI-IDENTIFIER CHECKS PASSED CLEANLY!")

    finally:
        db.query(OfferModel).filter(OfferModel.id == f"off-test-viewer-{test_id}").delete(synchronize_session=False)
        db.query(InterviewModel).filter(InterviewModel.id == f"int-test-viewer-{test_id}").delete(synchronize_session=False)
        db.query(ApplicationModel).filter(ApplicationModel.id == f"app-test-viewer-{test_id}").delete(synchronize_session=False)
        db.query(CandidateModel).filter(CandidateModel.id.in_([cand_id, f"cand-test-viewer-{test_id}"])).delete(synchronize_session=False)
        db.commit()
        db.close()
        print("Cleanup completed.")

if __name__ == "__main__":
    run_test()
